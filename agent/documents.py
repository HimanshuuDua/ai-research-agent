import csv
import io
import json
import re
from dataclasses import dataclass
from html.parser import HTMLParser

MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_DOCUMENT_CHARS = 120_000

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".html",
    ".htm",
    ".rtf",
}


@dataclass
class ExtractedDocument:
    filename: str
    text: str
    char_count: int
    truncated: bool
    format: str


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self._parts.append(data.strip())

    def get_text(self) -> str:
        return "\n".join(self._parts)


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _truncate(text: str, limit: int = MAX_DOCUMENT_CHARS) -> tuple[str, bool]:
    if len(text) <= limit:
        return text, False
    return text[:limit] + "\n\n[Document truncated due to size limit.]", True


def _extension(filename: str) -> str:
    dot = filename.rfind(".")
    if dot == -1:
        return ""
    return filename[dot:].lower()


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_rtf(data: bytes) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError as exc:
        raise ValueError("RTF support is unavailable.") from exc
    return rtf_to_text(data.decode("utf-8", errors="replace"))


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = [" | ".join(row) for row in reader if any(cell.strip() for cell in row)]
    return "\n".join(rows)


def _extract_json(data: bytes) -> str:
    payload = json.loads(data.decode("utf-8", errors="replace"))
    return json.dumps(payload, indent=2, ensure_ascii=False)


def _extract_html(data: bytes) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(data.decode("utf-8", errors="replace"))
    return parser.get_text()


def extract_document(data: bytes, filename: str) -> ExtractedDocument:
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError(f"File exceeds {MAX_UPLOAD_BYTES // (1024 * 1024)} MB limit.")

    ext = _extension(filename)
    if ext == ".doc":
        raise ValueError(
            "Legacy .doc files are not supported. Save as .docx or PDF and upload again."
        )
    if ext and ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {supported}")

    if ext == ".pdf":
        raw = _extract_pdf(data)
        doc_format = "pdf"
    elif ext == ".docx":
        raw = _extract_docx(data)
        doc_format = "docx"
    elif ext == ".rtf":
        raw = _extract_rtf(data)
        doc_format = "rtf"
    elif ext == ".csv":
        raw = _extract_csv(data)
        doc_format = "csv"
    elif ext == ".json":
        raw = _extract_json(data)
        doc_format = "json"
    elif ext in {".html", ".htm"}:
        raw = _extract_html(data)
        doc_format = "html"
    else:
        raw = data.decode("utf-8", errors="replace")
        doc_format = ext.lstrip(".") or "text"

    text = _normalize_whitespace(raw)
    if not text:
        raise ValueError("No readable text found in this document.")

    text, truncated = _truncate(text)
    return ExtractedDocument(
        filename=filename,
        text=text,
        char_count=len(text),
        truncated=truncated,
        format=doc_format,
    )
